from argparse import ArgumentParser
import pandas as pd
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn



# https://stackoverflow.com/questions/39006878/python-docx-add-horizontal-line
def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_feedback_column(mistake_row, document, style):
    feedback = mistake_row['Feedback']
    if str(feedback) != 'nan' and str(feedback) != 'v':
        document.add_paragraph(style=style).add_run("Our feedback:")
        document.add_paragraph(style=style).add_run(feedback)
        document.add_paragraph(style=style).add_run("")


def compile_feedback_for_instance(items, document):
    # sentence,sentence_id,target_idx,instance_id,target,
    # worker_id, general_comment, question_template_worker, question_worker, answer_worker, answer_range_worker, property_worker, comment_worker,
    # expert_id,                  question_template_expert, question_expert, answer_expert, answer_range_expert, property_expert, comment_expert, mismatch_type
    calibri_no_space = document.styles['Calibri No Spacing']
    sentence = items[0]['sentence']
    tokens = sentence.split()
    target_idx = int(items[0]['target_idx'])
    your_par = document.add_paragraph()
    your_par.style = "List Paragraph"
    tnr_name = "Times New Roman"
    tnr = document.styles['List Bullet - Times New Roman']


    your_par.add_run(" ".join(tokens[:target_idx])).font.name = tnr_name
    target_run = your_par.add_run(" " + tokens[target_idx] + " ")
    target_run.bold = True
    target_run.font.name = tnr_name
    your_par.add_run(" ".join(tokens[target_idx+1:])).font.name=tnr_name

    missed = [item for item in items if item['answer_expert'] and not item['answer_worker']]
    if missed:
        missed_par = document.add_paragraph(style=calibri_no_space)
        missed_par.add_run("You ")
        missed_word = missed_par.add_run("missed ")
        missed_word.font.color.rgb = RGBColor(255, 0, 0)
        missed_par.add_run("the following question-answer pairs: ")
        for miss in missed:
            missed_qa = f"{miss['question_expert']} --- {miss['answer_expert']}"
            document.add_paragraph(style=tnr).add_run(missed_qa)
            add_feedback_column(miss, document, calibri_no_space)

    error = [item for item in items if not item['answer_expert'] and item['answer_worker']]
    if error:
        err_p = document.add_paragraph(style=calibri_no_space)
        err_p.add_run("You have ")
        err_word = err_p.add_run("erroneously added ")
        err_word.font.color.rgb = RGBColor(255,0,0)
        # err_word.bold = True
        err_p.add_run("the following question-answer pairs: ")
        for err in error:
            err_qa = f"{err['question_worker']} --- {err['answer_worker']}"
            document.add_paragraph(err_qa, style=tnr)
            add_feedback_column(err, document, calibri_no_space)

    qs_mismatch = [item for item in items if item['answer_expert'] and item['answer_worker']]
    qs_mismatch = [item for item in qs_mismatch if item['question_expert'] != item['question_worker']]
    if qs_mismatch:
        q_paragraph = document.add_paragraph(style=calibri_no_space)
        q_paragraph.add_run("You have a ")
        q_word = q_paragraph.add_run("question mismatch ")
        q_word.font.color.rgb = RGBColor(255, 0, 0)
        q_paragraph.add_run("with the expert in the following question-answer pairs: ")
        for q in qs_mismatch:
            q_mismatch = f"Your QA: {q['question_worker']} --- {q['answer_worker']}\n" \
                         f"Expert's QA: {q['question_expert']} --- {q['answer_expert']}"
            document.add_paragraph(q_mismatch, style=tnr)
            add_feedback_column(q, document, calibri_no_space)

    insertHR(document.add_paragraph())


def compile_feedback(worker_id, worker_df):
    document = Document()
    styles = document.styles
    calibri = styles.add_style('Calibri Font', WD_STYLE_TYPE.PARAGRAPH)
    calibri.base_style = styles['Normal']
    calibri.font.name = "Calibri"

    calibri_no_space = styles.add_style('Calibri No Spacing', WD_STYLE_TYPE.PARAGRAPH)
    calibri_no_space.base_style = styles['No Spacing']
    calibri_no_space.font.name = "Calibri"

    tnr = styles.add_style("List Bullet - Times New Roman", WD_STYLE_TYPE.PARAGRAPH)
    tnr.base_style = styles['List Bullet']
    tnr.font.name = "Times New Roman"


    # TODO: add task title.
    document.add_heading('Feedback for your recent HITs', 0)
    p = document.add_paragraph(
        "This document is to help you understand "
        "some problems we have encountered while evaluating your recent work. "
        "Please read through these examples carefully to help you achieve the desired annotation "
        "requirements for this task. Workers that would complete successfully this short training round "
        "will be qualified to work on large batches of this task."
    )
    p.style = calibri
    insertHR(p)
    # p.font.name = "Calibri"
    # document.add_paragraph("Each of the next sections will show you an example of your recent work with some annotations that you have missed, ")

    for key, instance_df in worker_df.groupby(['sentence_id', 'target_idx']):
        instance_items = instance_df.to_dict(orient="records")
        compile_feedback_for_instance(instance_items, document)
    return document


def save_feedback(path_to_save, feedback):
    feedback.save(path_to_save)


def main(args):
    df = pd.read_csv(args.in_path)
    for c in ['answer', 'answer_range', 'comment', 'question', 'question_template']:
        for suffix in ("_worker", "_expert"):
            if c+suffix in df.columns:
                df[c+suffix].fillna("", inplace=True)
    for worker_id, worker_df in df.groupby('worker_id'):
        out_path = os.path.join(args.out_dir, f"feedback_{worker_id}.docx")
        feedback = compile_feedback(worker_id, worker_df)
        save_feedback(out_path, feedback)


if __name__ == "__main__":
    ap = ArgumentParser()
    # ap.add_argument("--in_path", default="annot_vs_expert.csv")
    # ap.add_argument("--out_dir", default="../training/feedback")
    ap.add_argument("--in_path", default="../../qanoun_share/with_feedback_column/worker_mistakes_batch1.csv")
    ap.add_argument("--out_dir", default="../../qanoun_share/with_feedback_column")
    main(ap.parse_args())